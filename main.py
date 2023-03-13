import requests
from bs4 import BeautifulSoup as BS
import csv
from docx import Document


def filter(x):
    for el in x:
        for word in el:
            if (len(word)) <= 3:
                el.remove(word)
    for el in x:
        if len(el) >= 2:
            el.append(' '.join(el))
            el.remove(el[0])
            el.remove(el[0])


def getStat():
    tds = []

    r = requests.get('https://fbref.com/en/comps/1/World-Cup-Stats')
    html = BS(r.text, 'html.parser')

    tr = html.find_all('tr')
    for td in tr:
        tds0 = td.find_all('td')
        for el in tds0:
            if el is not None:
                tds.append(el.text)

    teamNames = ['Netherlands', 'Senegal', 'Ecuador', 'Qatar', 'England', 'United States', 'Iran', 'Wales',
                 'Argentina', 'Poland', 'Mexico', 'Saudi Arabia', 'France', 'Australia', 'Tunisia', 'Denmark',
                 'Japan', 'Spain', 'Germany', 'Costa Rica', 'Morocco', 'Croatia', 'Belgium', 'Canada',
                 'Brazil', 'Switzerland', 'Cameroon', 'Serbia', 'Portugal', 'Korea Republic', 'Uruguay',
                 'Ghana']

    wins = tds[2::14]
    draws = tds[3::14]
    losses = tds[4::14]
    scored = tds[5::14]
    missed = tds[6::14]
    GD = tds[7::14]
    points = tds[8::14]
    xG = tds[9::14]
    xGA = tds[10::14]
    xGD = tds[11::14]
    xGD90 = tds[12::14]

    with open('data.csv', 'w') as file:
        writer = csv.writer(file)
        heads = ['Country', 'Wins', 'Draws', 'Losses', 'Scored', 'Missed', 'Differance', 'Points', 'xG', 'xGA', 'xGD',
                 'xGD/90']
        writer.writerow(heads)
        for w in range(len(teamNames)):
            writer.writerow(
                [teamNames[w], wins[w], draws[w], losses[w], scored[w], missed[w], GD[w], points[w], xG[w], xGA[w],
                 xGD[w], xGD90[w]])


getStat()

def getMatches():
    global winners, squad, opponent
    doc = Document('matches.docx')
    table = doc.tables[0]
    names = []
    winners = []

    for row in table.rows:
        string = ''
        check = False
        for cell in row.cells:
            string = cell.text
            names.append(string)
            rc = cell.paragraphs[0].runs[0]
            if rc.font.bold:
                winners.append(cell.text)
                check = True
        if check == False:
            winners.append('main draw')




    winners = [el.split(' ') for el in winners]
    squad = names[::2]

    squad = [el.split(' ') for el in squad]

    opponent = names[1::2]
    opponent = [el.split(' ') for el in opponent]


    filter(squad)
    squad = [elem[0] for elem in squad]
    squad.insert(0, 'Команда')

    filter(opponent)
    opponent = [elem[0] for elem in opponent]
    opponent.insert(0, 'Соперник')

    filter(winners)
    winners = [elem[0] for elem in winners]
    winners.insert(0, 'Победитель')

    with open('data1.csv', 'w') as file:
        writer = csv.writer(file)
        for w in range(len(squad)):
            writer.writerow([squad[w], opponent[w], winners[w]])

getMatches()